import uuid
import io
from typing import List
from sqlalchemy.orm import Session
from PyPDF2 import PdfReader
from docx import Document
from app import models

class DocumentService:
    def __init__(self):
        pass

    async def process_document(self, filename: str, file_type: str, file_size: int, content: bytes, user_id: str, db: Session):
        # Extract text
        text = self.extract_text(content, file_type)
        
        # Create document record
        doc_id = str(uuid.uuid4())
        document = models.Document(
            id=doc_id,
            filename=filename,
            file_type=file_type,
            file_size=file_size,
            content=text,
            user_id=user_id
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Chunk text
        chunks = self.chunk_text(text)
        
        for i, chunk in enumerate(chunks):
            chunk_record = models.DocumentChunk(
                document_id=doc_id,
                chunk_index=i,
                text_content=chunk,
                embedding=None
            )
            db.add(chunk_record)
        
        db.commit()
        document.processed_at = db.query(models.Document).filter(models.Document.id == doc_id).first().processed_at
        return document

    def extract_text(self, content: bytes, file_type: str) -> str:
        if file_type == "text/plain":
            return content.decode("utf-8")
        elif file_type == "application/pdf":
            pdf_reader = PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(io.BytesIO(content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        else:
            raise ValueError("Unsupported file type")

    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        words = text.split()
        chunks = []
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            chunks.append(chunk)
        return chunks